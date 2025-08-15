import pandas as pd
from newspaper import Article
import datetime
import subprocess
import csv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from bs4 import BeautifulSoup
import logging
import os
import re
import certifi
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
import urllib3
from webdriver_manager.chrome import ChromeDriverManager
import newspaper.network
import time
import ssl
from collections import defaultdict
from mots_listes import MOTS_CLES, SOCIETES, INTERMEDIAIRES
import json
from selenium.webdriver.chrome.service import Service
from datetime import datetime, timedelta   # üî• pour les dates
import sys
import os

temp_path = r'C:\Users\hp\AppData\Local\Temp\.newspaper_scraper\article_resources'
os.makedirs(temp_path, exist_ok=True)

# ‚úÖ Si on tourne depuis l'ex√©cutable PyInstaller, redirige newspaper vers les ressources packag√©es
if getattr(sys, 'frozen', False):  
    newspaper.settings.DATA_DIR = os.path.join(sys._MEIPASS, "newspaper", "resources")
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
ssl._create_default_https_context = ssl._create_unverified_context

# ===================== PATCH API CALENDRIER =====================
COUNTRY_CODE = "TN"  # üî• Tunisie, change si besoin (FR, US...)

def get_public_holidays(year, country_code=COUNTRY_CODE):
    """R√©cup√®re les jours f√©ri√©s pour une ann√©e donn√©e via nager.date (pas besoin de cl√© API)."""
    url = f"https://date.nager.at/api/v3/PublicHolidays/{year}/{country_code}"
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        holidays = response.json()
        return {datetime.strptime(h['date'], "%Y-%m-%d").date() for h in holidays}
    except Exception as e:
        logging.error(f"Erreur r√©cup√©ration jours f√©ri√©s : {e}")
        return set()

def calculer_dates_a_scraper():
    """Calcule dynamiquement les dates √† scraper, incluant les jours f√©ri√©s et week-ends selon le contexte."""
    today = datetime.today().date()
    holidays = get_public_holidays(today.year)

    dates_a_scraper = []

    # Commencer par hier
    yesterday = today - timedelta(days=1)
    jour = yesterday

    # Recule jusqu'au dernier jour ouvr√© avant le f√©ri√© ou week-end
    while True:
        dates_a_scraper.insert(0, jour)  # on ajoute au d√©but pour garder l'ordre chronologique
        if jour.weekday() < 5 and jour not in holidays:  # jour ouvr√© non f√©ri√©
            break  # on a atteint le dernier jour ouvr√©
        jour -= timedelta(days=1)

    # Si lundi, ajouter tous les jours du week-end + vendredi
    if today.weekday() == 0:  # lundi
        for i in range(1, 4):  # vendredi, samedi, dimanche
            d = today - timedelta(days=i)
            if d not in dates_a_scraper:
                dates_a_scraper.insert(0, d)

    logging.info(f"üìÖ Dates √† scraper : {dates_a_scraper}")
    return set(dates_a_scraper)


# ===================== PATCH NEWSPAPER SSL =====================
def patch_newspaper_ssl():
    session = requests.Session()
    session.verify = False
    newspaper.network._session = session

patch_newspaper_ssl()

# ===================== LOGGING =====================
# G√©n√©rer un nom de fichier log avec date et heure
timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
log_filename = f"agent_revue_presse_{timestamp}.log"

logging.basicConfig(
    filename=log_filename,
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# ===================== SCRAPING SELENIUM =====================
def parser_manuellement_selenium(url):
    options = Options()
    options.add_argument("--headless=new")
    options.add_argument("--disable-gpu")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--log-level=3")
    options.add_experimental_option('excludeSwitches', ['enable-logging'])

    service = Service(ChromeDriverManager().install(), log_path=os.devnull)
    driver = webdriver.Chrome(service=service, options=options)

    try:
        driver.get(url)
        WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.TAG_NAME, "article"))
        )
        html = driver.page_source
        soup = BeautifulSoup(html, 'html.parser')

        titre = soup.title.text.strip() if soup.title else "Titre indisponible"
        contenu = "\n".join(p.text.strip() for p in soup.find_all('p') if len(p.text.strip()) > 30)

        if len(contenu) < 100:
            raise ValueError("Contenu trop court pour √™tre un article.")

        return {
            "titre": titre,
            "contenu": contenu,
            "date": datetime.now(),
            "source": url,
            "lien": url
        }
    except Exception as e:
        logging.warning(f"√âchec du parsing Selenium pour {url} : {e}")
        return None
    finally:
        driver.quit()

# ===================== TELECHARGEMENT ARTICLE =====================
def download_article_with_retry(url, retries=2, delay=5):
    for attempt in range(1, retries + 1):
        try:
            art = Article(url)
            art.download()
            art.parse()
            if not art.text or len(art.text) < 100:
                raise ValueError("Contenu trop court")
            return art
        except Exception as e:
            logging.warning(f"Erreur parsing (tentative {attempt}) : {e}")
            time.sleep(delay)
    return None

# ===================== CHARGEMENT  CSV =====================
def charger_listes_depuis_csv(chemin_csv):
    societes = []
    intermediaires = []
    mots_cles = []
    try:
        with open(chemin_csv, encoding='utf-8-sig', newline='') as f:
            reader = csv.DictReader(f)
            for row in reader:
                if 'SOCIETES' in row and row['SOCIETES'].strip():
                    societes.append(row['SOCIETES'].strip())
                if 'INTERMEDIAIRES' in row and row['INTERMEDIAIRES'].strip():
                    intermediaires.append(row['INTERMEDIAIRES'].strip())
                if 'MOTS_CLES' in row and row['MOTS_CLES'].strip():
                    mots_cles.append(row['MOTS_CLES'].strip())

        if not (societes or intermediaires or mots_cles):
            raise ValueError("Les listes dans le CSV sont vides.")
        logging.info(f"Listes charg√©es depuis CSV : {len(societes)} soci√©t√©s, {len(intermediaires)} interm√©diaires, {len(mots_cles)} mots-cl√©s.")
        return societes, intermediaires, mots_cles
    except Exception as e:
        logging.warning(f"√âchec chargement listes CSV '{chemin_csv}': {e}")
        return None, None, None
# ===================== SCRAPER ARTICLES =====================

def scraper_articles(url_site, categorie=None):
    """
    Scrape un site web et retourne les articles trouv√©s contenant au moins
    un mot de MOTS_CLES, SOCIETES ou INTERMEDIAIRES.
    """
    articles_trouves = []
    nb_articles_selenium = 0

    headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:118.0) Gecko/20100101 Firefox/118.0',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
    'Accept-Language': 'fr-FR,fr;q=0.5',
    'Connection': 'keep-alive',
}


    def contient(cible, contenu):
        return any(mot.lower() in contenu for mot in cible)

    try:
        response = requests.get(url_site, headers=headers, timeout=60, verify = r"C:\Users\hp\Desktop\stage d'ete bourse\cacert.pem"
)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        liens = set()
        for a_tag in soup.find_all('a', href=True):
            href = a_tag['href']
            text = a_tag.get_text(separator=' ', strip=True).lower()

            # Normaliser les URL relatives en URL absolues
            if href.startswith('/'):
                href = requests.compat.urljoin(url_site, href)

            contenu = f"{href.lower()} {text}"

            # ‚úÖ V√©rifier la pr√©sence d‚Äôau moins un mot des listes
            if not (contient(MOTS_CLES, contenu) or contient(SOCIETES, contenu) or contient(INTERMEDIAIRES, contenu)):
                continue

            # ‚ùå Filtrer les liens ind√©sirables
            exclusions = ['/ads/', '/banner-', '.js', '.css', '.ico', '.svg', 'cdn-cgi', 'stbfinance.com.tn']
            if any(excl in href for excl in exclusions):
                continue

            # ‚ùå Filtrer certains formats de fichiers
            if any(href.lower().endswith(ext) for ext in ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.zip', '.rar']):
                continue

            # ‚úÖ V√©rifier URL valide
            if href.startswith(('http://', 'https://')) and not any(bad in href for bad in ['ck.php', 'INSERT_RANDOM_NUMBER_HERE']):
                liens.add(href)

        logging.info(f"{len(liens)} liens pertinents trouv√©s sur {url_site}")

        # ‚úÖ Parser les articles (limit√© √† 50 pour √©viter surcharge)
        for lien in list(liens)[:50]:
            if re.search(r'/category/|/categorie/', lien.lower()):
                continue
            # if "twitter.com" in lien.lower() or "x.com" in lien.lower():
            #     logging.info(f"‚ùå Ignored Twitter link: {lien}")
            #     continue

            try:
                resp = requests.get(lien, headers=headers, timeout=20)
                resp.raise_for_status()

                art = Article(lien)
                art.set_html(resp.text)
                art.parse()

                if art.text and len(art.text) > 100 and art.title and len(art.title) > 10:
                    date_article = art.publish_date or datetime.now()
                    logging.debug(f"Extraction article: '{art.title}', date: {date_article}")

                    articles_trouves.append({
                        "titre": art.title,
                        "contenu": art.text,
                        "date": date_article,
                        "source": getattr(art, 'source_url', url_site),
                        "lien": lien,
                        "categorie": categorie
                    })

                    logging.debug(f"‚úÖ Article extrait : {art.title} de {lien}")

            except Exception as e:
                logging.warning(f"‚ö†Ô∏è Article non pars√© : {lien} -> {e}")

    except requests.RequestException as e:
        logging.error(f"‚ùå Erreur requ√™te HTTP : {url_site} -> {e}")
    except Exception as e:
        logging.error(f"‚ùå Erreur scraping de {url_site} : {e}")

    if nb_articles_selenium > 0:
        logging.info(f"{nb_articles_selenium} articles r√©cup√©r√©s via Selenium pour {url_site}")

    return articles_trouves, nb_articles_selenium



def scraper_indices_boursiers(urls_indices):
    def extraire_pourcentage(val):
        if not val:
            return "NC"
        match = re.search(r'-?\d+[\.,]?\d*%?', val)
        if match:
            return match.group().replace(',', '.').strip('%') + '%'
        return "NC"

    all_indices = []

    for url in urls_indices:
        logging.info(f"Scraping indices boursiers depuis {url}")
        try:
            headers = {'User-Agent': 'Mozilla/5.0'}
            r = requests.get(url, headers=headers, timeout=20)
            r.raise_for_status()
            soup = BeautifulSoup(r.text, 'html.parser')
            
            if "bvmt.com.tn" in url:
                rows = soup.select('table.table-condensed tbody tr')
                for row in rows:
                    cells = row.find_all('td')
                    if len(cells) >= 6:
                        nom = cells[0].get_text(strip=True)
                        # üîÑ Inverser ici : cells[5] = Var. 2025 et cells[3] = Var. du jour
                        var_jour = extraire_pourcentage(cells[5].get_text(strip=True))  
                        var_annee = extraire_pourcentage(cells[3].get_text(strip=True))  
                        all_indices.append((nom, var_jour, var_annee))

            elif "www.boursorama.com/bourse/indices/internationaux" in url:
                rows = soup.select('table.c-table tbody tr')
                for row in rows:
                    cells = row.find_all('td')
                    if len(cells) >= 4:
                        nom = cells[0].get_text(strip=True)
                        var_jour = extraire_pourcentage(cells[3].get_text(strip=True))  
                        var_annee = extraire_pourcentage(cells[2].get_text(strip=True))  
                        all_indices.append((nom, var_jour, var_annee))

            elif "investing.com/indices" in url:
                rows = soup.select('table.genTbl.closedTbl.elpTbl.elp20.tblIndices tbody tr')
                for row in rows:
                    cells = row.find_all('td')
                    if len(cells) >= 8:
                        nom = cells[1].get_text(strip=True)
                        var_jour = extraire_pourcentage(cells[4].get_text(strip=True))
                        var_annee = "NC"
                        all_indices.append((nom, var_jour, var_annee))

            elif "www.egx.com.eg/ar/EGX_Error.aspx?aspxerrorpath=/ar/Indices.aspx" in url:
                rows = soup.select('table#ctl00_PlaceHolderMain_gvIndex tr')
                for row in rows[1:]:
                    cells = row.find_all('td')
                    if len(cells) >= 4:
                        nom = cells[0].get_text(strip=True)
                        var_jour = extraire_pourcentage(cells[2].get_text(strip=True))
                        var_annee = "NC"
                        all_indices.append((nom, var_jour, var_annee))
                        
            elif "www.casablanca-bourse.com/fr" in url:
                rows = soup.select('table.table.table-striped tbody tr')
                for row in rows:
                    cells = row.find_all('td')
                    if len(cells) >= 4:
                        nom = cells[0].get_text(strip=True)
                        var_jour = extraire_pourcentage(cells[3].get_text(strip=True))  
                        var_annee = extraire_pourcentage(cells[2].get_text(strip=True))  
                        all_indices.append((nom, var_jour, var_annee))

            elif "countryeconomy.com/stock-exchange" in url:
                rows = soup.select('table.table-hover tbody tr')
                for row in rows:
                    cells = row.find_all('td')
                    if len(cells) >= 5:
                        nom = cells[0].get_text(strip=True)
                        var_jour = extraire_pourcentage(cells[4].get_text(strip=True))  
                        var_annee = extraire_pourcentage(cells[2].get_text(strip=True))  
                        all_indices.append((nom, var_jour, var_annee))

            else:
                logging.warning(f"Site non pris en charge pour indices : {url}")

        except Exception as e:
            logging.error(f"Erreur scraping indices {url} : {e}")

    logging.info(f"Total indices boursiers r√©cup√©r√©s : {len(all_indices)}")

    def nettoyer_nom(nom):
        return nom.lower().replace(' ', '').replace('-', '')

    indices_a_garder = ["Tunindex", "MASI", "Egypte", "Tadawul", "CAC", "DAX", "FTSE", "Dow Jones"]
    indices_a_garder_clean = [nettoyer_nom(i) for i in indices_a_garder]

    filtered_indices = []
    for indice in all_indices:
        nom_clean = nettoyer_nom(indice[0])
        if any(i in nom_clean for i in indices_a_garder_clean):
            filtered_indices.append((indice[0], indice[1], indice[2]))

    logging.info(f"Indices boursiers filtr√©s : {len(filtered_indices)}")
    return filtered_indices

# --- Filtrage articles ---

def contient_un_mot(texte, liste_mots):
    """Retourne True si au moins un mot de liste_mots est trouv√© dans texte."""
    if not liste_mots:
        return False
    mots_echappes = [re.escape(mot) for mot in liste_mots if mot.strip()]
    if not mots_echappes:
        return False
    pattern = r'\b(?:' + '|'.join(mots_echappes) + r')\b'
    return bool(re.search(pattern, texte, re.IGNORECASE))

def filtrer_article(article):
    """
    Retient l'article s'il contient AU MOINS UN mot d'une des listes.
    """
    texte = f"{article.get('titre', '')} {article.get('contenu', '')}".lower()

    mot_cle_ok = contient_un_mot(texte, MOTS_CLES)
    societe_ok = contient_un_mot(texte, SOCIETES)
    intermediaire_ok = contient_un_mot(texte, INTERMEDIAIRES)

    logging.debug(f"[FILTRAGE] '{article.get('titre', 'Sans titre')}' -> "
                  f"Mots-cl√©s={mot_cle_ok} | Soci√©t√©s={societe_ok} | Interm√©diaires={intermediaire_ok}")

    return mot_cle_ok or societe_ok or intermediaire_ok

# --- R√©sum√© avec Ollama ---
def resumer_article(contenu):
    contenu_limite = contenu[:4000]
    prompt = (
    "Tu es un assistant expert en finance charg√© de r√©sumer des articles de presse √©conomiques ou boursiers.\n"
    "Ton r√©sum√© doit √™tre r√©dig√© exclusivement en fran√ßais, avec un ton neutre, clair et professionnel.\n"
    "Le r√©sum√© doit √™tre informatif et contextualis√©, m√™me si le contenu est limit√©.\n"
    "Reprends exactement tous les noms propres, acronymes, indices boursiers, entreprises, institutions et interm√©diaires mentionn√©s dans le texte, sans les reformuler ni les remplacer par des synonymes.\n"
    "Reprends √©galement tous les chiffres, pourcentages, dates ou p√©riodes tels qu‚Äôils apparaissent dans l'article.\n"
    "Ne te contente pas de citer les entit√©s : pr√©cise leur r√¥le ou leur lien avec les faits d√©crits.\n"
    "Si l‚Äôarticle contient peu d‚Äôinformations ou seulement des noms sans contexte, indique clairement que l'article manque de d√©tails.\n"
    "R√©sume le texte suivant en 5 √† 7 phrases maximum, en mettant en √©vidence :\n"
    "- les faits cl√©s\n"
    "- les donn√©es chiffr√©es importantes\n"
    "- les acteurs √©conomiques cit√©s et leur r√¥le\n"
    "- les dates ou p√©riodes importantes\n\n"
    "Texte √† r√©sumer :\n"
    f"{contenu_limite}\n\n"
    "R√©sum√© :"
)


    try:
        cmd = ['ollama', 'run', 'mistral', prompt]
        logging.info(f"R√©sum√© IA Ollama pour contenu {len(contenu_limite)} caract√®res.")
        resultat = subprocess.run(cmd, capture_output=True, text=True, check=True, timeout=360, encoding='utf-8')
        resume = resultat.stdout.strip()
        if not resume:
            raise ValueError("R√©sum√© vide retourn√© par Ollama.")
        return resume
    except FileNotFoundError:
        logging.error("Ollama non trouv√©.")
        return "R√©sum√© indisponible : Ollama non trouv√©."
    except subprocess.CalledProcessError as e:
        logging.error(f"Erreur Ollama (code {e.returncode}): {e.stderr}")
        return "R√©sum√© indisponible : Erreur Ollama."
    except subprocess.TimeoutExpired:
        logging.error("Timeout Ollama.")
        return "R√©sum√© indisponible : Timeout Ollama."
    except Exception as e:
        logging.error(f"Erreur r√©sum√© IA : {e}")
        return "R√©sum√© indisponible."

# --- Ajout lien hypertexte dans docx ---
def ajouter_hyperlien(paragraph, url, texte):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = texte
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)
    return hyperlink

# --- G√©n√©ration du document Word avec regroupement par cat√©gorie ---
def generer_revue_presse(articles, indices_boursiers_data=None, fichier_sortie=None):
    if fichier_sortie is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        fichier_sortie = f"revue_presse_{timestamp}.docx"
        doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    heading = doc.add_heading("Revue de presse", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_heading = heading.runs[0]
    run_heading.font.size = Pt(24)

    date_para = doc.add_paragraph(f"Du {datetime.today().strftime('%d/%m/%Y')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.runs[0]
    date_run.font.size = Pt(14)
    doc.add_paragraph()

    if indices_boursiers_data:
        doc.add_heading("Indices Boursiers", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Indice'
        hdr_cells[1].text = 'Var. du jour'
        hdr_cells[2].text = 'Var. 2025'

        for i, hdr in enumerate(hdr_cells):
            for run in hdr.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(11)
            hdr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for nom, var_jour, var_annee in indices_boursiers_data:
            row_cells = table.add_row().cells
            row_cells[0].text = nom
            row_cells[1].text = var_jour
            row_cells[2].text = var_annee
            for cell in row_cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph("-" * 80)
        doc.add_paragraph()

    for article in articles:
        if isinstance(article['date'], datetime) and article['date'].tzinfo is not None:
            article['date'] = article['date'].replace(tzinfo=None)

    articles_tries = sorted(
        articles,
        key=lambda x: x['date'] if isinstance(x['date'], datetime) else datetime.min,
        reverse=True
    )

    articles_par_categorie = {}
    for art in articles_tries:
        cat = art.get("categorie", "Inconnu")
        articles_par_categorie.setdefault(cat, []).append(art)

    for categorie, articles_cat in articles_par_categorie.items():
        doc.add_heading(categorie, level=1)

        for art in articles_cat:
            p_titre = doc.add_paragraph()
            run_titre = p_titre.add_run(art['titre'])
            run_titre.bold = True
            run_titre.font.size = Pt(12)

            date_display = art['date'].strftime('%d/%m/%Y') if isinstance(art['date'], datetime) else str(art['date'])
            p_info = doc.add_paragraph()
            run_info = p_info.add_run(f"Source: {art['source']} ‚Äì {date_display}")
            run_info.italic = True
            run_info.font.size = Pt(10)

            p_link = doc.add_paragraph()
            ajouter_hyperlien(p_link, art['lien'], art['lien'])

            if p_link.runs:
                for run in p_link.runs:
                    run.font.size = Pt(10)

            p_resume = doc.add_paragraph(art.get('resume', 'R√©sum√© non disponible'))
            p_resume.runs[0].font.size = Pt(11)

            doc.add_paragraph("-" * 80)
        doc.add_paragraph()

    try:
        doc.save(fichier_sortie)
        logging.info(f"Document Word sauvegard√© : {fichier_sortie}")
    except Exception as e:
        logging.error(f"Erreur sauvegarde document Word : {e}")


# --- D√©finition des sites par cat√©gories ---
def obtenir_sites_par_categorie(fichier_csv='Liste des sites √† consulter.csv'):
    # Liste par d√©faut au cas o√π le CSV est indisponible ou vide
    liste_par_defaut = {
        "M√©dias Nationaux": [
            "https://www.ilboursa.com/",
            "http://www.tustex.com/",
            "https://www.webmanagercenter.com/",
            "https://africanmanager.com/categorie/finances/",
            "https://africanmanager.com/categorie/actualites/",
            "https://africanmanager.com/categorie/la-une/",
            "https://www.businessnews.com.tn/dernieres-news",
            "https://www.businessnews.com.tn/Dossiers",
            "https://www.leconomistemaghrebin.com/",
            "https://www.espacemanager.com/",
            "https://lapresse.tn/",
            "https://www.entreprises-magazine.com/",
            "https://universnews.tn/",
            "https://www.millim.tn/",
            "https://irbe7.com/",
            "https://www.challenges.tn/category/economie/bourse/",
            "https://radioexpressfm.com/fr/",
            "https://www.leaders.com.tn/",
            "https://www.tap.info.tn/fr/"
        ],
        "M√©dias Internationaux": [
            "https://www.lesechos.fr/finance-marches",
            "https://www.lesechos.fr/bourse",
            "https://www.lesechos.fr/economie-france",
            "https://www.boursorama.com/bourse/actualites/marches/",
            "https://www.boursorama.com/bourse/actualites/finances/",
            "https://www.latribune.fr/",
            "https://www.boursorama.com/bourse/actualites/",
            "https://notreafrik.com/",
            "https://www.jeuneafrique.com/"
        ],
        "Indices Boursiers": [
            "https://www.bvmt.com.tn",
            "https://www.boursorama.com/bourse/indices/internationaux",
            "https://www.investing.com/indices",
            "https://www.egx.com.eg/ar/EGX_Error.aspx?aspxerrorpath=/ar/Indices.aspx",
            "https://www.casablanca-bourse.com/fr",
            "https://countryeconomy.com/stock-exchange"
        ]
    }

    if not os.path.isfile(fichier_csv):
        print(f"‚ö†Ô∏è Fichier {fichier_csv} non trouv√©, utilisation de la liste par d√©faut.")
        return liste_par_defaut

    sites = defaultdict(list)
    try:
        with open(fichier_csv, newline='', encoding='utf-8') as csvfile:
            lecteur = csv.DictReader(csvfile)
            rows = list(lecteur)
            if not rows:
                print(f"‚ö†Ô∏è Fichier {fichier_csv} vide, utilisation de la liste par d√©faut.")
                return liste_par_defaut
            for ligne in rows:
                categorie = ligne.get('categorie')
                url = ligne.get('url')
                if categorie and url:
                    sites[categorie.strip()].append(url.strip())
                else:
                    print("‚ö†Ô∏è Ligne CSV mal form√©e, utilisation partielle des donn√©es.")
        # Si aucune donn√©e valide extraite, fallback sur liste par d√©faut
        if not sites:
            print("‚ö†Ô∏è Aucune donn√©e valide extraite du CSV, utilisation de la liste par d√©faut.")
            return liste_par_defaut

        return dict(sites)
    except Exception as e:
        print(f"‚ö†Ô∏è Erreur lors de la lecture du fichier CSV : {e}\nUtilisation de la liste par d√©faut.")
        return liste_par_defaut
CHECKPOINT_FILE = "checkpoint8.json"

def charger_checkpoint(path=CHECKPOINT_FILE):
    if os.path.exists(path):
        try:
            with open(path, 'r', encoding='utf-8') as f:
                contenu = f.read().strip()
                if not contenu:
                    raise ValueError("Fichier checkpoint vide")
                checkpoint = json.loads(contenu)
                logging.info("Checkpoint charg√©.")
                return checkpoint
        except Exception as e:
            logging.error(f"Erreur chargement checkpoint: {e}")

    checkpoint_defaut = {
        "sites_termine": [],
        "articles_traite": [],
        "articles_resume": []
    }
    sauvegarder_checkpoint(checkpoint_defaut, path)
    logging.info("Nouveau checkpoint cr√©√©.")
    return checkpoint_defaut

def sauvegarder_checkpoint(checkpoint, path=CHECKPOINT_FILE):
    try:
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(checkpoint, f, ensure_ascii=False, indent=2)
        logging.info(f"Checkpoint sauvegard√© dans {path} avec {len(checkpoint.get('articles_resume', []))} articles.")
    except Exception as e:
        logging.error(f"Erreur sauvegarde checkpoint: {e}")


# ===================== AGENT PRINCIPAL =====================
def agent():
    global SOCIETES, INTERMEDIAIRES, MOTS_CLES  # <-- global en premier

    base_path = r"C:\Users\hp\Desktop\stage d'ete bourse"
    checkpoint_path = os.path.join(base_path, "checkpoint8.json")
    chemin_csv = os.path.join(base_path, "Liste SC & IB & Mots cl√©s.csv")

    logging.info("üöÄ D√©marrage de l'agent IA de revue de presse.")

    # Essayer de charger depuis CSV
    societes, intermediaires, mots_cles = charger_listes_depuis_csv(chemin_csv)

    # Si √©chec du chargement (retour None), fallback vers mots_listes.py
    if societes is None or intermediaires is None or mots_cles is None:
        logging.info("Chargement depuis CSV √©chou√©, fallback vers mots_listes.py")
        from mots_listes import SOCIETES, INTERMEDIAIRES, MOTS_CLES as mots_cles_defaut
        societes = SOCIETES
        intermediaires = INTERMEDIAIRES
        mots_cles = mots_cles_defaut

    # Mettre √† jour les variables globales
    SOCIETES = societes
    INTERMEDIAIRES = intermediaires
    MOTS_CLES = mots_cles

    # V√©rification
    if not (MOTS_CLES or SOCIETES or INTERMEDIAIRES):
        logging.error("‚ùå Les listes (mots-cl√©s, soci√©t√©s, interm√©diaires) sont vides apr√®s chargement.")
        return
    logging.info(f"‚úÖ {len(MOTS_CLES)} mots-cl√©s, {len(SOCIETES)} soci√©t√©s, {len(INTERMEDIAIRES)} interm√©diaires charg√©s.")


    # üåç 2Ô∏è‚É£ Charger la liste des sites
    sites_par_categorie = obtenir_sites_par_categorie("Liste des sites √† consulter.csv")
    urls_indices = sites_par_categorie.get("Indices Boursiers", [])

    # üìå 3Ô∏è‚É£ Charger checkpoint
    checkpoint_data = charger_checkpoint(checkpoint_path)
    articles_checkpoint = checkpoint_data.get("articles_resume", [])
    liens_deja_tries = set(a['lien'] for a in articles_checkpoint)
    sites_termine = set(checkpoint_data.get("sites_termine", []))
    logging.info(f"üìä Checkpoint : {len(sites_termine)} sites d√©j√† trait√©s, {len(liens_deja_tries)} articles d√©j√† r√©sum√©s.")

    # üìÖ 4Ô∏è‚É£ Calcul des dates √† scraper
    dates_a_scraper = calculer_dates_a_scraper()
    logging.info(f"üìÜ Dates cibl√©es pour scraping : {dates_a_scraper}")

    tous_articles = []
    total_selenium_articles = 0
    total_articles_traite = 0
    total_articles_resumes = 0
    total_articles_filtres = 0

    # üì∞ 5Ô∏è‚É£ Boucle sur les sites (hors indices)
    for categorie, urls in sites_par_categorie.items():
        if categorie == "Indices Boursiers":
            continue

        for site in urls:
            if site in sites_termine:
                logging.info(f"‚è≠Ô∏è Site d√©j√† trait√© : {site}")
                continue

            logging.info(f"üåê Scraping du site : {site} (cat√©gorie : {categorie})")
            # ‚úÖ plus besoin de passer les listes en argument
            articles, nb_selenium = scraper_articles(site, categorie=categorie)
            logging.info(f"üìë {len(articles)} articles bruts trouv√©s sur {site}")
            total_selenium_articles += nb_selenium

            # üéØ Filtrer par date
            articles_nouveaux = [
                a for a in articles
                if a['lien'] not in liens_deja_tries
                and isinstance(a['date'], datetime)
                and a['date'].date() in dates_a_scraper
            ]
            logging.info(f"üìÜ {len(articles_nouveaux)} articles dat√©s des jours cibl√©s.")

            # üè∑Ô∏è Filtrage sur mots-cl√©s / soci√©t√©s / interm√©diaires
            for art in articles_nouveaux:
                total_articles_traite += 1
                if filtrer_article(art):
                    logging.info(f"üìù R√©sum√© de l'article : {art['titre']}")
                    art['resume'] = resumer_article(art['contenu'])
                    if art['resume']:
                        tous_articles.append(art)
                        total_articles_resumes += 1
                        # ‚ûï Ajout au checkpoint
                        articles_checkpoint.append({
                            "lien": art['lien'],
                            "titre": art.get('titre', ''),
                            "resume": art.get('resume', ''),
                            "categorie": art.get('categorie', '')
                        })
                        liens_deja_tries.add(art['lien'])
                    else:
                        logging.warning(f"‚ö†Ô∏è R√©sum√© vide pour article : {art['titre']}")
                else:
                    total_articles_filtres += 1
                    logging.info(f"‚ùå Article exclu par filtre : {art['titre']}")

            # ‚úÖ Marquer le site comme termin√©
            sites_termine.add(site)

            # üíæ Sauvegarder checkpoint apr√®s chaque site
            checkpoint_final = {
                "sites_termine": list(sites_termine),
                "articles_traite": list(liens_deja_tries),
                "articles_resume": articles_checkpoint
            }
            sauvegarder_checkpoint(checkpoint_final, checkpoint_path)

    # üìä R√©sum√© scraping
    logging.info(f"üìä STATISTIQUES SCRAPING : {total_articles_traite} trait√©s, {total_articles_resumes} r√©sum√©s, {total_articles_filtres} filtr√©s.")

    # üìà 6Ô∏è‚É£ Scraper les indices boursiers
    logging.info("üìà Scraping des indices boursiers...")
    indices_boursiers_data = scraper_indices_boursiers(urls_indices)
    if indices_boursiers_data:
        logging.info(f"‚úÖ {len(indices_boursiers_data)} indices r√©cup√©r√©s.")
    else:
        logging.warning("‚ö†Ô∏è Aucun indice boursier r√©cup√©r√©.")
        indices_boursiers_data = None

    # üìÑ 7Ô∏è‚É£ G√©n√©ration du document Word
    logging.info("üìÑ G√©n√©ration du document Word de la revue de presse...")
    generer_revue_presse(tous_articles, indices_boursiers_data)

    logging.info("‚úÖ Agent termin√© avec succ√®s !")


# ---- Ex√©cution principale ----
if __name__ == "__main__":
    print("Bienvenue dans l'agent IA de revue de presse.")
    choix = input("Tape 'go' pour lancer l'agent, ou autre pour quitter : ")
    if choix.lower() == 'go':
        agent()
    else:
        print("Sortie du programme.")
